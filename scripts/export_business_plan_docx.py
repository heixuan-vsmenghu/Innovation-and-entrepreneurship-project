from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "business-plan" / "business_plan_draft.md"
OUTPUT = ROOT / "deliverables" / "创策云UniPlanAI商业计划书.docx"


def set_run_font(run, size=12, bold=False, font_name="宋体"):
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def set_paragraph_font(paragraph, size=12, bold=False, font_name="宋体"):
    for run in paragraph.runs:
        set_run_font(run, size=size, bold=bold, font_name=font_name)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if bold else WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text.strip())
    set_run_font(run, size=10.5, bold=bold)


def set_table_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "6")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "4B5563")
        borders.append(border)
    tbl_pr.append(borders)


def clean_inline(text):
    text = text.replace("`", "")
    text = text.replace("**", "")
    return text.strip()


def is_table_start(lines, index):
    if index + 1 >= len(lines):
        return False
    return lines[index].lstrip().startswith("|") and re.match(r"^\s*\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?\s*$", lines[index + 1])


def parse_table(lines, index):
    table_lines = []
    while index < len(lines) and lines[index].lstrip().startswith("|"):
        table_lines.append(lines[index])
        index += 1
    rows = []
    for raw in table_lines:
        cells = [clean_inline(cell) for cell in raw.strip().strip("|").split("|")]
        rows.append(cells)
    if len(rows) >= 2:
        rows.pop(1)
    return rows, index


def add_markdown_table(document, rows):
    if not rows:
        return
    col_count = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"
    set_table_borders(table)
    for row_index, row in enumerate(rows):
        for col_index in range(col_count):
            text = row[col_index] if col_index < len(row) else ""
            set_cell_text(table.cell(row_index, col_index), text, bold=row_index == 0)
    document.add_paragraph()


def add_body_paragraph(document, text, is_bullet=False):
    style = "List Bullet" if is_bullet else None
    paragraph = document.add_paragraph(style=style)
    if not is_bullet:
        paragraph.paragraph_format.first_line_indent = Cm(0.74)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.25
    run = paragraph.add_run(clean_inline(text))
    set_run_font(run)


def configure_document(document):
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

    styles = document.styles
    for style_name in ("Normal", "Body Text"):
        style = styles[style_name]
        style.font.name = "宋体"
        style.font.size = Pt(12)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    for style_name, size in (("Heading 1", 18), ("Heading 2", 15), ("Heading 3", 13)):
        style = styles[style_name]
        style.font.name = "宋体"
        style.font.size = Pt(size)
        style.font.bold = True
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def export_docx():
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    text = SOURCE.read_text(encoding="utf-8")
    lines = text.splitlines()
    document = Document()
    configure_document(document)

    index = 0
    while index < len(lines):
        raw = lines[index]
        line = raw.strip()
        if not line:
            index += 1
            continue

        if is_table_start(lines, index):
            rows, index = parse_table(lines, index)
            add_markdown_table(document, rows)
            continue

        if line.startswith("# "):
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(clean_inline(line[2:]))
            set_run_font(run, size=20, bold=True, font_name="宋体")
            index += 1
            continue

        if line.startswith("## "):
            title = clean_inline(line[3:])
            if title.startswith("二、目录") or title.startswith("三、执行概要"):
                document.add_page_break()
            paragraph = document.add_heading(title, level=1)
            set_paragraph_font(paragraph, size=15, bold=True)
            index += 1
            continue

        if line.startswith("### "):
            paragraph = document.add_heading(clean_inline(line[4:]), level=2)
            set_paragraph_font(paragraph, size=13, bold=True)
            index += 1
            continue

        if line.startswith("- "):
            add_body_paragraph(document, line[2:], is_bullet=True)
            index += 1
            continue

        add_body_paragraph(document, line)
        index += 1

    document.save(OUTPUT)
    print(f"exported {OUTPUT}")


if __name__ == "__main__":
    export_docx()
